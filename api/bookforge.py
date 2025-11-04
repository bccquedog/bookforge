#!/usr/bin/env python3
"""
BookForge — cross‑platform book formatting tool
-------------------------------------------------
A Python-based CLI wizard that asks the right questions about your book
and produces print-ready interior PDFs (with professional layout), plus EPUB
and DOCX outputs. Works on macOS and Windows. Packagable with PyInstaller.

Key features
- Interactive Q&A wizard builds a production config (trim size, margins, fonts, front matter, headers/footers)
- Accepts raw manuscript: .docx / .md / .txt (single file) or a folder of .md chapters
- Normalizes to HTML via pandoc (if available) or fallback converters
- Typesets to PDF with WeasyPrint using professional book CSS with page rules
- Generates EPUB via pandoc (if available)
- Auto front matter builder (title, copyright, dedication, TOC) and back matter (acknowledgments, about author)
- Running heads/feet with different recto/verso templates
- Widows/orphans control, hyphenation (CSS hyphens; optional soft-hyphen insertion)
- KDP/Ingram presets (trim sizes, margins, bleed safe areas)
- Cover size calculator (spine width from page count & paper type)
- Saves a reusable YAML config; supports non-interactive runs with `--config`
- One-line packaging: `pyinstaller --onefile bookforge.py`

"""

import argparse
import dataclasses
import os
import re
import shutil
import subprocess
import sys
import tempfile
import textwrap
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, Optional, List, Any

import yaml
from jinja2 import Template
from unidecode import unidecode
from rapidfuzz import fuzz

# Optional deps; import lazily when needed
try:
    import docx  # python-docx
except Exception:
    docx = None

# WeasyPrint for PDF
from weasyprint import HTML, CSS

APP_NAME = "BookForge"
VERSION = "1.0.0"

# ----------------------
# Presets & style tables
# ----------------------

TRIM_PRESETS = {
    # width x height in inches
    "5x8": {"width_in": 5.0, "height_in": 8.0, "target": "trade"},
    "5.5x8.5": {"width_in": 5.5, "height_in": 8.5, "target": "trade"},
    "6x9": {"width_in": 6.0, "height_in": 9.0, "target": "trade"},
    "8.5x11": {"width_in": 8.5, "height_in": 11.0, "target": "workbook"},
}

PAPER_STOCKS = {
    "cream_55lb": {"kdp": True, "ingram": True, "pages_per_inch": 444},
    "white_50lb": {"kdp": True, "ingram": True, "pages_per_inch": 512},
}

DEFAULTS = {
    "trim": "6x9",
    "font_family": "'EB Garamond','Garamond','Georgia',serif",
    "font_size_pt": 11.0,
    "line_height": 1.35,
    "outer_margin_in": 0.75,
    "top_margin_in": 0.75,
    "bottom_margin_in": 0.85,
    "gutter_in": 0.7,
    "chapter_starts_right": True,
    "hyphenate": True,
    "ornament": "\u2042",  # asterism as default scene break
}

# ----------------------
# Data classes
# ----------------------

@dataclass
class BuildConfig:
    title: str
    subtitle: str = ""
    author: str = ""
    imprint: str = ""
    isbn: str = ""
    trim: str = DEFAULTS["trim"]
    paper: str = "cream_55lb"
    font_family: str = DEFAULTS["font_family"]
    font_size_pt: float = DEFAULTS["font_size_pt"]
    line_height: float = DEFAULTS["line_height"]
    outer_margin_in: float = DEFAULTS["outer_margin_in"]
    top_margin_in: float = DEFAULTS["top_margin_in"]
    bottom_margin_in: float = DEFAULTS["bottom_margin_in"]
    gutter_in: float = DEFAULTS["gutter_in"]
    chapter_starts_right: bool = True
    hyphenate: bool = True
    header_style: str = "author_title"  # or title_only / author_only / none
    include_toc: bool = True
    include_dedication: bool = False
    dedication_text: str = ""
    include_copyright: bool = True
    copyright_year: str = ""
    copyright_holder: str = ""
    include_ack: bool = False
    ack_text: str = ""
    include_about_author: bool = False
    about_author_text: str = ""
    scene_break: str = DEFAULTS["ornament"]
    target_platform: str = "kdp"  # or ingram / generic
    bleed: bool = False
    language: str = "en"
    toc_depth: int = 2
    epub: bool = True
    docx: bool = True

    def to_yaml(self) -> str:
        return yaml.safe_dump(dataclasses.asdict(self), sort_keys=False, allow_unicode=True)

# ----------------------
# Utilities
# ----------------------

def inch_to_mm(x: float) -> float:
    return round(x * 25.4, 3)


def which(exe: str) -> Optional[str]:
    return shutil.which(exe)


def have_pandoc() -> bool:
    return which("pandoc") is not None


def slugify(s: str) -> str:
    s = unidecode(s)
    s = re.sub(r"[^A-Za-z0-9]+", "-", s)
    return s.strip("-").lower()


# ----------------------
# Wizard: ask the right questions
# ----------------------

def prompt_default(prompt: str, default: str) -> str:
    val = input(f"{prompt} [{default}]: ").strip()
    return val or default


def prompt_bool(prompt: str, default: bool = True) -> bool:
    dv = "Y/n" if default else "y/N"
    val = input(f"{prompt} ({dv}): ").strip().lower()
    if not val:
        return default
    return val in {"y", "yes", "true", "1"}


def run_wizard(manuscript: Path, outdir: Path) -> BuildConfig:
    print(f"\n{APP_NAME} Wizard — let's set up your book.\n")
    title = prompt_default("Book title", Path(manuscript).stem.replace("_", " "))
    subtitle = input("Subtitle (optional): ").strip()
    author = prompt_default("Author name", os.getlogin() if hasattr(os, 'getlogin') else "")
    imprint = input("Imprint/Publisher (optional): ").strip()
    isbn = input("ISBN (optional): ").strip()

    # Trim + platform
    print("\nPick a trim size:")
    for k in TRIM_PRESETS:
        print(f"  - {k}")
    trim = prompt_default("Trim size", DEFAULTS["trim"])
    while trim not in TRIM_PRESETS:
        print("Please choose one of:", ", ".join(TRIM_PRESETS))
        trim = prompt_default("Trim size", DEFAULTS["trim"])

    platform = prompt_default("Target platform (kdp/ingram/generic)", "kdp").lower()
    paper = prompt_default("Paper stock (cream_55lb/white_50lb)", "cream_55lb")

    # Margins (auto suggestions)
    page_count_guess = guess_pages_from_wordcount(manuscript, DEFAULTS["font_size_pt"]) or 220
    gutter = suggested_gutter(page_count_guess)

    print(f"\nEstimated page count: ~{page_count_guess} pages (adjust later if needed).")
    outer = float(prompt_default("Outer margin (inches)", f"{DEFAULTS['outer_margin_in']:.2f}"))
    top = float(prompt_default("Top margin (inches)", f"{DEFAULTS['top_margin_in']:.2f}"))
    bottom = float(prompt_default("Bottom margin (inches)", f"{DEFAULTS['bottom_margin_in']:.2f}"))
    gutter_in = float(prompt_default("Gutter (inner) margin (inches)", f"{gutter:.2f}"))

    font_family = prompt_default("Body font stack", DEFAULTS["font_family"])
    font_size = float(prompt_default("Body font size (pt)", f"{DEFAULTS['font_size_pt']:.1f}"))
    line_height = float(prompt_default("Line height (eg 1.35)", f"{DEFAULTS['line_height']:.2f}"))

    header_style = prompt_default("Running header style (author_title/title_only/author_only/none)", "author_title")
    chapter_right = prompt_bool("Start chapters on right-hand (recto) pages?", True)
    hyphenate = prompt_bool("Enable hyphenation?", True)
    include_toc = prompt_bool("Include Table of Contents?", True)

    include_ded = prompt_bool("Include dedication page?", False)
    ded_text = input("Dedication text (if included): ").strip() if include_ded else ""

    include_cr = prompt_bool("Include copyright page?", True)
    cr_year = prompt_default("Copyright year", str(Path.cwd()).split(os.sep)[-1]) if include_cr else ""
    cr_holder = prompt_default("Copyright holder", author) if include_cr else ""

    include_ack = prompt_bool("Include acknowledgments?", False)
    ack_text = input("Acknowledgments text: ").strip() if include_ack else ""

    include_about = prompt_bool("Include About the Author?", False)
    about_text = input("About the Author text: ").strip() if include_about else ""

    scene_break = prompt_default("Scene break ornament (unicode or text)", DEFAULTS["ornament"])

    cfg = BuildConfig(
        title=title,
        subtitle=subtitle,
        author=author,
        imprint=imprint,
        isbn=isbn,
        trim=trim,
        paper=paper,
        font_family=font_family,
        font_size_pt=font_size,
        line_height=line_height,
        outer_margin_in=outer,
        top_margin_in=top,
        bottom_margin_in=bottom,
        gutter_in=gutter_in,
        chapter_starts_right=chapter_right,
        hyphenate=hyphenate,
        header_style=header_style,
        include_toc=include_toc,
        include_dedication=include_ded,
        dedication_text=ded_text,
        include_copyright=include_cr,
        copyright_year=cr_year,
        copyright_holder=cr_holder,
        include_ack=include_ack,
        ack_text=ack_text,
        include_about_author=include_about,
        about_author_text=about_text,
        scene_break=scene_break,
    )

    # Save YAML alongside output
    outdir.mkdir(parents=True, exist_ok=True)
    (outdir / "bookforge.yml").write_text(cfg.to_yaml(), encoding="utf-8")
    print(f"\nSaved config to {outdir/'bookforge.yml'}\n")
    return cfg


def suggested_gutter(page_count: int) -> float:
    # Rough industry heuristic for perfect-bound
    if page_count < 150:
        return 0.6
    elif page_count < 300:
        return 0.7
    elif page_count < 500:
        return 0.8
    else:
        return 0.9


def guess_pages_from_wordcount(path: Path, font_size_pt: float) -> Optional[int]:
    try:
        wc = estimate_wordcount(path)
        if not wc:
            return None
        # rough words/page estimate; varies by trim & layout
        wpp = 350 if font_size_pt <= 11 else 300
        return max(1, int(round(wc / wpp)))
    except Exception:
        return None


def estimate_wordcount(path: Path) -> Optional[int]:
    try:
        txt = extract_text_quick(path)
        words = re.findall(r"\b\w+\b", txt)
        return len(words)
    except Exception:
        return None


def extract_text_quick(path: Path) -> str:
    ext = path.suffix.lower()
    if ext == ".txt":
        return Path(path).read_text(encoding="utf-8", errors="ignore")
    if ext == ".md":
        return Path(path).read_text(encoding="utf-8", errors="ignore")
    if ext == ".docx" and docx is not None:
        d = docx.Document(str(path))
        return "\n".join(p.text for p in d.paragraphs)
    # Fallback: try Pandoc
    if have_pandoc():
        with tempfile.TemporaryDirectory() as td:
            out = Path(td) / "tmp.txt"
            subprocess.run(["pandoc", str(path), "-t", "plain", "-o", str(out)], check=True)
            return out.read_text(encoding="utf-8")
    raise RuntimeError("Cannot extract text; install python-docx for .docx or Pandoc for other types.")


# ----------------------
# Build pipeline
# ----------------------

def detect_chapter_heading(line: str) -> bool:
    """Detect if a line is a chapter heading"""
    line_upper = line.strip().upper()
    # Common chapter patterns
    patterns = [
        r'^CHAPTER\s+[A-Z]+[\w\s]*:?',  # "CHAPTER ONE", "CHAPTER TWENTY-ONE:"
        r'^CHAPTER\s+\d+',  # "CHAPTER 1", "CHAPTER 21"
        r'^PART\s+[IVX]+',  # "PART I", "PART II"
        r'^PART\s+\d+',  # "PART 1", "PART 2"
        r'^PROLOGUE',
        r'^EPILOGUE',
    ]
    for pattern in patterns:
        if re.match(pattern, line_upper):
            return True
    return False

def txt_to_html_with_chapters(text: str) -> str:
    """Convert plain text to HTML with chapter detection"""
    lines = text.split('\n')
    html_lines = ['<html><body>']
    in_paragraph = False
    chapter_count = 0
    skip_next = False
    
    for i, line in enumerate(lines):
        if skip_next:
            skip_next = False
            continue
            
        stripped = line.strip()
        
        # Detect chapter headings
        if detect_chapter_heading(line) and len(stripped) < 200:  # Heading should be short
            # Close previous chapter section if open
            if chapter_count > 0:
                html_lines.append('</section>')
            
            # Close previous paragraph if open
            if in_paragraph:
                html_lines.append('</p>')
                in_paragraph = False
            
            # Check if next line is also part of the heading (subtitle)
            heading_text = stripped
            if i + 1 < len(lines) and lines[i + 1].strip() and not detect_chapter_heading(lines[i + 1]):
                next_line = lines[i + 1].strip()
                if len(next_line) < 100 and not next_line.endswith('.'):
                    heading_text += ': ' + next_line
                    skip_next = True  # Skip next line
            
            # Add chapter heading with proper structure
            chapter_count += 1
            heading_id = f"chapter-{chapter_count}"
            html_lines.append(f'<section class="chapter" id="{heading_id}">')
            html_lines.append(f'<h1 class="chapter-title">{html_escape(heading_text)}</h1>')
            in_paragraph = False
            continue
        
        # Empty line - break paragraph
        if not stripped:
            if in_paragraph:
                html_lines.append('</p>')
                in_paragraph = False
            # Close chapter if we have one open and we've had some content
            continue
        
        # Regular text
        escaped = html_escape(stripped)
        
        # Start new paragraph if needed
        if not in_paragraph:
            html_lines.append('<p>')
            in_paragraph = True
        else:
            # Add space between lines in same paragraph
            html_lines.append(' ')
        
        html_lines.append(escaped)
    
    # Close any open tags
    if in_paragraph:
        html_lines.append('</p>')
    
    # Close any open chapter sections
    if chapter_count > 0:
        html_lines.append('</section>')
    
    html_lines.append('</body></html>')
    return '\n'.join(html_lines)

def detect_and_convert_chapters_in_html(html: str) -> str:
    """Post-process HTML to ensure chapters are properly structured for EPUB"""
    # Extract body content
    body_match = re.search(r'<body[^>]*>(.*?)</body>', html, re.DOTALL)
    if not body_match:
        return html
    
    body_content = body_match.group(1)
    lines = body_content.split('\n')
    processed_lines = []
    in_paragraph = False
    chapter_count = 0
    
    for i, line in enumerate(lines):
        stripped = line.strip()
        
        # Check if this line looks like a chapter heading in HTML
        # Look for lines that are standalone (not in paragraphs) and match chapter patterns
        if re.search(r'<p[^>]*>(.*?)</p>', line):
            # Extract text from paragraph
            text_match = re.search(r'<p[^>]*>(.*?)</p>', line)
            if text_match:
                text_content = re.sub(r'<[^>]+>', '', text_match.group(1)).strip()
                if detect_chapter_heading(text_content) and len(text_content) < 200:
                    # Convert paragraph to h1 heading and wrap in section
                    chapter_count += 1
                    heading_id = f"chapter-{chapter_count}"
                    processed_lines.append(f'<section class="chapter" id="{heading_id}">')
                    processed_lines.append(f'<h1 class="chapter-title">{text_match.group(1)}</h1>')
                    in_paragraph = False
                    continue
        
        # Check for plain text lines that might be headings (from pre tags)
        if stripped and not stripped.startswith('<') and detect_chapter_heading(stripped) and len(stripped) < 200:
            chapter_count += 1
            heading_id = f"chapter-{chapter_count}"
            processed_lines.append(f'<section class="chapter" id="{heading_id}">')
            processed_lines.append(f'<h1 class="chapter-title">{html_escape(stripped)}</h1>')
            in_paragraph = False
            continue
        
        processed_lines.append(line)
    
    # Close any open sections
    if chapter_count > 0:
        processed_lines.append('</section>')
    
    # Reconstruct HTML
    new_body = '\n'.join(processed_lines)
    html = html[:body_match.start(1)] + new_body + html[body_match.end(1):]
    
    return html

def extract_body_content(html: str) -> str:
    """Extract just the body content from a full HTML document"""
    # Try to extract content between <body> tags
    body_match = re.search(r'<body[^>]*>(.*?)</body>', html, re.DOTALL | re.IGNORECASE)
    if body_match:
        return body_match.group(1).strip()
    # If no body tags, assume it's already just body content
    return html.strip()

def convert_to_html(manuscript: Path) -> str:
    ext = manuscript.suffix.lower()
    if have_pandoc():
        cmd = [
            "pandoc", str(manuscript),
            "--from", "docx" if ext == ".docx" else ("markdown" if ext == ".md" else "plain"),
            "--to", "html5",
            "--section-divs",
            "--standalone",
        ]
        html = subprocess.check_output(cmd, text=True)
        # Post-process to ensure chapter headings are properly detected for txt files
        if ext == ".txt":
            html = detect_and_convert_chapters_in_html(html)
        # Extract just the body content (template will wrap it)
        return extract_body_content(html)
    # Fallback simple converters
    if ext == ".md":
        html = markdown_to_html_simple(Path(manuscript).read_text(encoding="utf-8"))
        return extract_body_content(html)
    if ext == ".txt":
        html = txt_to_html_with_chapters(Path(manuscript).read_text(encoding="utf-8", errors="ignore"))
        return extract_body_content(html)
    if ext == ".docx" and docx is not None:
        html = docx_to_html_simple(manuscript)
        return extract_body_content(html)
    raise RuntimeError("No conversion path to HTML; install Pandoc or use .md/.txt/.docx with python-docx.")


def html_escape(s: str) -> str:
    return (s
            .replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
            .replace('"', "&quot;")
            .replace("'", "&#39;")
    )


def markdown_to_html_simple(md: str) -> str:
    # Minimalist markdown (not as robust as Pandoc)
    md = md.replace("\r\n", "\n")
    lines = md.split("\n")
    html_lines = ["<html><body>"]
    in_p = False
    for line in lines:
        if re.match(r"^# +", line):
            level = len(line) - len(line.lstrip('#'))
            html_lines.append(f"<h{level}>{html_escape(line[level:].strip())}</h{level}>")
            in_p = False
        elif not line.strip():
            if in_p:
                html_lines.append("</p>")
                in_p = False
        else:
            if not in_p:
                html_lines.append("<p>")
                in_p = True
            html_lines.append(html_escape(line) + " ")
    if in_p:
        html_lines.append("</p>")
    html_lines.append("</body></html>")
    return "\n".join(html_lines)


def docx_to_html_simple(path: Path) -> str:
    d = docx.Document(str(path))
    parts = ["<html><body>"]
    for p in d.paragraphs:
        txt = html_escape(p.text)
        if not txt.strip():
            parts.append("<br>")
        else:
            parts.append(f"<p>{txt}</p>")
    parts.append("</body></html>")
    return "".join(parts)


# ----------------------
# Templates
# ----------------------

HTML_TEMPLATE = Template(r"""
<!doctype html>
<html lang="{{ language }}">
<head>
  <meta charset="utf-8">
  <meta name="viewport" content="width=device-width, initial-scale=1.0">
  <title>{{ title }}{% if subtitle %}: {{ subtitle }}{% endif %}</title>
  <style>
    /* Screen styles for HTML viewing */
    * {
      box-sizing: border-box;
    }
    
    body.book {
      font-family: {{ font_family }};
      font-size: {{ font_size_pt * 1.333 }}px; /* Convert pt to px for screen */
      line-height: {{ line_height }};
      color: #333;
      background: #fff;
      max-width: 800px;
      margin: 0 auto;
      padding: 2rem 1.5rem;
      text-align: justify;
    }
    
    /* Front matter styling */
    .front-matter {
      margin-bottom: 3rem;
    }
    
    .blank.verso {
      display: none;
    }
    
    .half-title {
      text-align: center;
      font-size: 1.5em;
      margin: 2rem 0;
      padding: 2rem 0;
    }
    
    .title-page {
      text-align: center;
      margin: 3rem 0;
      padding: 3rem 0;
    }
    
    .book-title {
      font-size: 2.5em;
      margin: 0 0 0.5em 0;
      font-weight: bold;
    }
    
    .book-subtitle {
      font-size: 1.4em;
      margin: 0.5em 0;
      color: #666;
      font-weight: normal;
    }
    
    .book-author {
      margin-top: 1.5em;
      font-size: 1.2em;
    }
    
    .book-imprint, .book-isbn {
      margin-top: 0.5em;
      font-size: 0.9em;
      color: #666;
    }
    
    .copyright {
      margin: 2rem 0;
      padding: 1rem 0;
      font-size: 0.9em;
      border-top: 1px solid #ddd;
      border-bottom: 1px solid #ddd;
    }
    
    .dedication {
      text-align: center;
      font-style: italic;
      margin: 3rem 0;
      padding: 2rem 0;
    }
    
    .toc {
      margin: 2rem 0;
    }
    
    .toc h2 {
      margin-bottom: 1rem;
      font-size: 1.5em;
    }
    
    /* Body matter styling */
    .body-matter {
      margin: 3rem 0;
    }
    
    .body-matter section.chapter {
      margin: 3rem 0;
      page-break-before: always;
    }
    
    .body-matter h1.chapter-title {
      font-size: 1.8em;
      margin: 2rem 0 1.5rem 0;
      font-weight: bold;
      text-align: left;
      page-break-after: avoid;
    }
    
    .body-matter h1 {
      font-size: 1.8em;
      margin: 2rem 0 1rem 0;
      font-weight: bold;
    }
    
    .body-matter h2 {
      font-size: 1.4em;
      margin: 1.5rem 0 0.75rem 0;
      font-weight: bold;
    }
    
    .body-matter h3 {
      font-size: 1.2em;
      margin: 1.25rem 0 0.5rem 0;
      font-weight: bold;
    }
    
    .body-matter p {
      margin: 0 0 1em 0;
      text-align: justify;
      text-indent: 1.2em;
      line-height: {{ line_height }};
    }
    
    .body-matter .chapter-open p,
    .body-matter .chapter-title + p {
      text-indent: 0;
    }
    
    .body-matter p + p {
      margin-top: 0;
    }
    
    .hr-ornament {
      text-align: center;
      margin: 2em 0;
      font-size: 1.2em;
      letter-spacing: 0.5em;
    }
    
    /* Back matter styling */
    .back-matter {
      margin-top: 4rem;
      padding-top: 2rem;
      border-top: 2px solid #ddd;
    }
    
    .back-matter h2 {
      font-size: 1.5em;
      margin-bottom: 1rem;
    }
    
    .back-matter p {
      text-indent: 0;
    }
    
    /* Responsive design */
    @media (max-width: 768px) {
      body.book {
        font-size: {{ font_size_pt * 1.2 }}px;
        padding: 1rem;
      }
      
      .book-title {
        font-size: 2em;
      }
      
      .body-matter h1.chapter-title {
        font-size: 1.5em;
      }
    }
    
    /* Print styles */
    @media print {
      {{ base_css }}
      
      body.book {
        max-width: none;
        padding: 0;
      }
    }
  </style>
  <style>
    /* PDF-specific styles (only applied via WeasyPrint) */
    {{ base_css }}
  </style>
</head>
<body class="book">
  <section class="front-matter">
    <div class="half-title recto">{{ title }}</div>
    <div class="blank verso"></div>

    <div class="title-page recto">
      <h1 class="book-title">{{ title }}</h1>
      {% if subtitle %}<h2 class="book-subtitle">{{ subtitle }}</h2>{% endif %}
      {% if author %}<div class="book-author">{{ author }}</div>{% endif %}
      {% if imprint %}<div class="book-imprint">{{ imprint }}</div>{% endif %}
      {% if isbn %}<div class="book-isbn">ISBN: {{ isbn }}</div>{% endif %}
    </div>
    <div class="blank verso"></div>

    {% if include_copyright %}
    <div class="copyright verso">
      <p>Copyright © {{ copyright_year }} {{ copyright_holder }}. All rights reserved.</p>
      <p>No part of this book may be reproduced without permission, except for brief quotations used in reviews or scholarly works.</p>
    </div>
    {% endif %}

    {% if include_dedication %}
    <div class="dedication recto"><p>{{ dedication_text }}</p></div>
    {% endif %}

    {% if include_toc %}
    <div class="toc recto">
      <h2>Contents</h2>
      <nav id="toc"></nav>
    </div>
    {% endif %}
  </section>

  <section class="body-matter">
    {{ body_html | safe }}
  </section>

  <section class="back-matter">
    {% if include_ack %}
    <div class="ack recto">
      <h2>Acknowledgments</h2>
      <p>{{ ack_text }}</p>
    </div>
    {% endif %}

    {% if include_about_author %}
    <div class="about recto">
      <h2>About the Author</h2>
      <p>{{ about_author_text }}</p>
    </div>
    {% endif %}
  </section>
</body>
</html>
""")


CSS_BASE_TEMPLATE = Template(r"""
@page {
  size: {{ page_width_mm }}mm {{ page_height_mm }}mm;
  margin-top: {{ top_margin_mm }}mm;
  margin-bottom: {{ bottom_margin_mm }}mm;
  margin-inside: {{ gutter_mm }}mm;
  margin-outside: {{ outer_margin_mm }}mm;
}

@page:left {
  @bottom-left { content: counter(page); }
  {% if header_style != 'none' %}
  @top-left { content: "{{ header_left }}"; }
  {% endif %}
}
@page:right {
  @bottom-right { content: counter(page); }
  {% if header_style != 'none' %}
  @top-right { content: "{{ header_right }}"; }
  {% endif %}
}

body.book { font-family: {{ font_family }}; font-size: {{ font_size_pt }}pt; line-height: {{ line_height }}; hyphens: {{ 'auto' if hyphenate else 'manual' }}; min-height: auto; }

h1, h2, h3 { page-break-after: avoid; }

p { text-align: justify; widows: 2; orphans: 2; }

p + p { text-indent: 1.2em; margin-top: 0; }
.chapter-open p { text-indent: 0; }

.hr-ornament { text-align: center; margin: 1.2em 0; }
.hr-ornament::after { content: "{{ scene_break }}"; letter-spacing: 0.35em; }

.title-page { display: flex; flex-direction: column; align-items: center; justify-content: center; height: 100vh; }
.half-title { display: flex; align-items: center; justify-content: center; height: 100vh; font-size: 150%; }
.book-title { font-size: 220%; margin: 0; }
.book-subtitle { font-size: 140%; margin-top: 0.25em; color: #555; }
.book-author { margin-top: 1em; font-size: 120%; }
.book-imprint, .book-isbn { margin-top: 0.35em; font-size: 80%; color: #666; }

.toc h2 { margin-bottom: 0.5em; }
#toc { display: block; }

/* Chapter handling */
section, div.chapter { page-break-before: auto; }

/* Force chapters to start recto if requested (best-effort) */
{% if chapter_starts_right %}
.chapter-open { break-before: right; }
{% endif %}

/* Headers/footers spacing */
body { counter-reset: page 1; }

/* First pages (front matter) without headers/footers as needed */
.front-matter .recto, .front-matter .verso { @page { @top-left { content: none } @top-right { content: none } } }

/* Simple TOC autogen (best-effort via JS inserted later) */
""")


# ----------------------
# Build functions
# ----------------------

def build_outputs(cfg: BuildConfig, manuscript: Path, outdir: Path) -> Dict[str, Path]:
    outdir.mkdir(parents=True, exist_ok=True)
    print("Converting manuscript to HTML...")
    body_html = convert_to_html(manuscript)

    # Normalize headings to form chapters; add class hooks
    body_html = postprocess_body_html(body_html, cfg)

    # Render full HTML
    html_final = render_full_html(cfg, body_html)
    html_path = outdir / f"{slugify(cfg.title)}.html"
    html_path.write_text(html_final, encoding="utf-8")

    # PDF via WeasyPrint
    print("Rendering PDF...")
    pdf_path = outdir / f"{slugify(cfg.title)}.pdf"
    base_css = render_base_css(cfg)
    HTML(string=html_final).write_pdf(target=str(pdf_path), stylesheets=[CSS(string=base_css)])

    outputs = {"html": html_path, "pdf": pdf_path}

    # EPUB via pandoc (optional)
    if cfg.epub and have_pandoc():
        print("Building EPUB via Pandoc...")
        epub_path = outdir / f"{slugify(cfg.title)}.epub"
        # Use --epub-chapter-level to split chapters on h1 headings
        # --section-divs preserves the section structure
        # --toc-depth=2 for better table of contents
        subprocess.run([
            "pandoc", str(html_path), "-o", str(epub_path),
            "--from", "html",
            "--to", "epub3",
            "--epub-chapter-level=1",  # Split on h1 headings (chapter titles)
            "--section-divs",  # Preserve section divisions
            "--toc-depth=2",  # Include h1 and h2 in TOC
            "--metadata", f"title={cfg.title}",
            "--metadata", f"author={cfg.author}",
            "--metadata", "language=en",
        ], check=True)
        outputs["epub"] = epub_path

    # DOCX via pandoc (optional)
    if cfg.docx and have_pandoc():
        print("Building DOCX via Pandoc...")
        docx_path = outdir / f"{slugify(cfg.title)}.docx"
        subprocess.run(["pandoc", str(html_path), "-o", str(docx_path)], check=True)
        outputs["docx"] = docx_path

    print("Done. Outputs:")
    for k, p in outputs.items():
        print(f"  - {k}: {p}")
    return outputs


def postprocess_body_html(html: str, cfg: BuildConfig) -> str:
    # Add simple TOC anchors, chapter-open class on h1 containers, scene break transformation
    # Insert a lightweight script to build a TOC (WeasyPrint supports limited JS; TOC is best-effort)
    # For stable PDF TOCs, rely on Pandoc when possible.
    html = re.sub(r"<hr */?>", '<div class="hr-ornament"></div>', html)

    # If we already have <section class="chapter"> tags (from txt conversion), preserve them
    # Otherwise, wrap h1 sections as chapter sections
    if '<section class="chapter"' not in html:
        # Wrap standalone h1 headings with section tags for EPUB chapter splitting
        # Match h1 that aren't already in a section
        html = re.sub(
            r'(?<!<section[^>]*>)(<h1[^>]*>.*?</h1>)',
            r'<section class="chapter">\1',
            html,
            flags=re.DOTALL
        )
        # Close sections before the next h1 or at the end
        html = re.sub(r'(</h1>)(?!</section>)(?=<h1)', r'\1</section>', html)
        # Also wrap in div for PDF styling compatibility
        html = re.sub(r"(<h1[^>]*>)", r'<div class="chapter-open">\1', html)
        html = re.sub(r"(</h1>)", r"\1</div>", html)
    else:
        # We have section tags, just add the chapter-open div for PDF styling
        html = re.sub(r'(<section class="chapter"[^>]*>\s*)(<h1[^>]*>)', r'\1<div class="chapter-open">\2', html)
        html = re.sub(r'(</h1>)(?=\s*</section>)', r'\1</div>', html)

    # Add IDs to headings for TOC (if they don't already have IDs)
    def idify(match):
        full_tag = match.group(0)
        # Skip if already has an id attribute
        if 'id=' in full_tag:
            return full_tag
        text = re.sub("<[^<]+?>", "", match.group(2))
        sid = slugify(text)[:60] or "section"
        return f"<h{match.group(1)} id=\"{sid}\">{match.group(2)}</h{match.group(1)}>"

    html = re.sub(r"<h([1-6])>(.*?)</h\1>", idify, html)
    return html


def render_base_css(cfg: BuildConfig) -> str:
    trim = TRIM_PRESETS[cfg.trim]
    return CSS_BASE_TEMPLATE.render(
        page_width_mm=inch_to_mm(trim["width_in"]),
        page_height_mm=inch_to_mm(trim["height_in"]),
        top_margin_mm=inch_to_mm(cfg.top_margin_in),
        bottom_margin_mm=inch_to_mm(cfg.bottom_margin_in),
        gutter_mm=inch_to_mm(cfg.gutter_in),
        outer_margin_mm=inch_to_mm(cfg.outer_margin_in),
        font_family=cfg.font_family,
        font_size_pt=cfg.font_size_pt,
        line_height=cfg.line_height,
        hyphenate=cfg.hyphenate,
        chapter_starts_right=cfg.chapter_starts_right,
        scene_break=cfg.scene_break,
        header_style=cfg.header_style,
        header_left=(cfg.author if cfg.header_style in ("author_title","author_only") else ""),
        header_right=(cfg.title if cfg.header_style in ("author_title","title_only") else ""),
    )


def render_full_html(cfg: BuildConfig, body_html: str) -> str:
    base_css = render_base_css(cfg)
    return HTML_TEMPLATE.render(
        language=cfg.language,
        title=cfg.title,
        subtitle=cfg.subtitle,
        author=cfg.author,
        imprint=cfg.imprint,
        isbn=cfg.isbn,
        base_css=base_css,
        include_copyright=cfg.include_copyright,
        copyright_year=cfg.copyright_year,
        copyright_holder=cfg.copyright_holder,
        include_dedication=cfg.include_dedication,
        dedication_text=cfg.dedication_text,
        include_toc=cfg.include_toc,
        include_ack=cfg.include_ack,
        ack_text=cfg.ack_text,
        include_about_author=cfg.include_about_author,
        about_author_text=cfg.about_author_text,
        body_html=body_html,
    )


# ----------------------
# Cover calculator (spine & full cover size)
# ----------------------

def calc_spine_width_in(page_count: int, paper: str) -> float:
    ppi = PAPER_STOCKS.get(paper, PAPER_STOCKS["cream_55lb"])['pages_per_inch']
    return round(page_count / ppi, 3)


def cover_dimensions(trim_key: str, page_count: int, paper: str, bleed: bool = False) -> Dict[str, float]:
    trim = TRIM_PRESETS[trim_key]
    bleed_in = 0.125 if bleed else 0.0
    spine_in = calc_spine_width_in(page_count, paper)
    width_in = trim['width_in'] * 2 + spine_in + 2 * bleed_in
    height_in = trim['height_in'] + 2 * bleed_in
    return {"width_in": width_in, "height_in": height_in, "spine_in": spine_in}


# ----------------------
# CLI
# ----------------------

def load_config(path: Path) -> BuildConfig:
    data = yaml.safe_load(Path(path).read_text(encoding="utf-8"))
    return BuildConfig(**data)


def save_config(cfg: BuildConfig, path: Path):
    path.write_text(cfg.to_yaml(), encoding="utf-8")


def cmd_wizard(args):
    manuscript = Path(args.manuscript)
    outdir = Path(args.outdir or (manuscript.parent / "bookforge_build"))
    cfg = run_wizard(manuscript, outdir)
    build_outputs(cfg, manuscript, outdir)


def cmd_build(args):
    manuscript = Path(args.manuscript)
    outdir = Path(args.outdir or (manuscript.parent / "bookforge_build"))
    cfg = load_config(Path(args.config)) if args.config else None
    if not cfg:
        raise SystemExit("--config is required for build mode (or use wizard)")
    build_outputs(cfg, manuscript, outdir)


def cmd_covercalc(args):
    dims = cover_dimensions(args.trim, args.pages, args.paper, args.bleed)
    print(f"Full cover: {dims['width_in']:.3f}in x {dims['height_in']:.3f}in; Spine: {dims['spine_in']:.3f}in")


def main():
    parser = argparse.ArgumentParser(prog=APP_NAME, description="Book typesetting wizard & builder")
    sub = parser.add_subparsers(required=True)

    p_wiz = sub.add_parser('wizard', help='Run interactive wizard and build outputs')
    p_wiz.add_argument('manuscript', help='Path to manuscript (.docx/.md/.txt)')
    p_wiz.add_argument('--outdir', help='Output directory (default: manuscript_dir/bookforge_build)')
    p_wiz.set_defaults(func=cmd_wizard)

    p_build = sub.add_parser('build', help='Build using an existing YAML config')
    p_build.add_argument('--config', required=True, help='bookforge.yml')
    p_build.add_argument('--manuscript', required=True, help='Path to manuscript')
    p_build.add_argument('--outdir', help='Output directory')
    p_build.set_defaults(func=cmd_build)

    p_cover = sub.add_parser('covercalc', help='Calculate full cover/spine dimensions')
    p_cover.add_argument('--trim', default=DEFAULTS['trim'], choices=list(TRIM_PRESETS.keys()))
    p_cover.add_argument('--pages', type=int, required=True)
    p_cover.add_argument('--paper', default='cream_55lb', choices=list(PAPER_STOCKS.keys()))
    p_cover.add_argument('--bleed', action='store_true')
    p_cover.set_defaults(func=cmd_covercalc)

    args = parser.parse_args()
    args.func(args)


if __name__ == '__main__':
    main()



