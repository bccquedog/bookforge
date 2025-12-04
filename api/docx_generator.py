"""
Professional DOCX Generator for BookForge
Generates properly formatted Word documents with professional book styling.

Uses python-docx for direct document creation with:
- Title pages and front matter
- Chapter headings with proper styles
- Body text with first-line indents
- Page breaks before chapters
- Headers/footers with page numbers
- Professional typography
"""

import re
import logging
from pathlib import Path
from typing import List, Tuple, Optional, Dict, Any
from dataclasses import dataclass
from enum import Enum

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, Inches, Twips, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not available. Professional DOCX generation disabled.")


class ContentType(Enum):
    """Types of content in a manuscript"""
    TITLE = "title"
    CHAPTER_HEADING = "chapter_heading"
    SECTION_HEADING = "section_heading"
    FRONT_MATTER_HEADING = "front_matter_heading"
    BACK_MATTER_HEADING = "back_matter_heading"
    PARAGRAPH = "paragraph"
    SCENE_BREAK = "scene_break"
    BLANK = "blank"


@dataclass
class ContentBlock:
    """A block of content with its type and text"""
    content_type: ContentType
    text: str
    level: int = 1  # Heading level (1 for chapter, 2 for section, etc.)


@dataclass
class DocxConfig:
    """Configuration for DOCX generation"""
    title: str
    author: str = ""
    subtitle: str = ""

    # Page setup
    page_width_inches: float = 6.0
    page_height_inches: float = 9.0
    top_margin_inches: float = 0.75
    bottom_margin_inches: float = 0.85
    left_margin_inches: float = 0.75  # outer margin
    right_margin_inches: float = 0.7  # gutter

    # Typography
    body_font: str = "Georgia"
    body_font_size_pt: float = 11.0
    heading_font: str = "Georgia"
    chapter_font_size_pt: float = 16.0
    title_font_size_pt: float = 24.0
    line_spacing: float = 1.35
    first_line_indent_inches: float = 0.25

    # Front/back matter
    include_title_page: bool = True
    include_copyright: bool = True
    copyright_year: str = ""
    copyright_holder: str = ""
    include_dedication: bool = False
    dedication_text: str = ""
    include_toc: bool = True
    include_about_author: bool = False
    about_author_text: str = ""

    # Scene breaks
    scene_break_symbol: str = "* * *"


# Chapter heading patterns (comprehensive)
CHAPTER_PATTERNS = [
    r'^CHAPTER\s+[A-Z]+[\w\s\-]*:',  # "CHAPTER ONE: Title" or "CHAPTER TWENTY-ONE:"
    r'^CHAPTER\s+\d+:',  # "CHAPTER 1: Title"
    r'^CHAPTER\s+[A-Z]+[\w\s\-]*$',  # "CHAPTER ONE" (end of line)
    r'^CHAPTER\s+\d+$',  # "CHAPTER 1" (end of line)
    r'^PART\s+[IVX]+',  # "PART I", "PART II"
    r'^PART\s+\d+',  # "PART 1", "PART 2"
    r'^PROLOGUE',
    r'^EPILOGUE',
    r'^INTERLUDE',
    r'^ACT\s+[IVX\d]+',  # "ACT I", "ACT 1"
]

# Front matter patterns
FRONT_MATTER_PATTERNS = [
    r'^DEDICATION$',
    r'^ACKNOWLEDGMENTS?$',
    r'^ACKNOWLEDGEMENTS?$',
    r'^PREFACE$',
    r'^FOREWORD$',
    r'^INTRODUCTION$',
    r'^NOTE\s+TO\s+READER',
    r'^AUTHOR[\'S]?\s+NOTE',
]

# Back matter patterns
BACK_MATTER_PATTERNS = [
    r'^ABOUT\s+THE\s+AUTHOR',
    r'^ALSO\s+BY',
    r'^BIBLIOGRAPHY',
    r'^GLOSSARY',
    r'^INDEX$',
    r'^NOTES$',
    r'^APPENDIX',
    r'^AFTERWORD',
    r'^BONUS',
]


def detect_content_type(line: str) -> Tuple[ContentType, int]:
    """
    Detect the type of content a line represents.

    Returns:
        Tuple of (ContentType, heading_level)
    """
    stripped = line.strip()
    if not stripped:
        return ContentType.BLANK, 0

    upper = stripped.upper()

    # Scene breaks (multiple asterisks, dashes, or other ornaments)
    if re.match(r'^[\*\-\~\#]{3,}$', stripped) or stripped in ['***', '* * *', '---', '~~~']:
        return ContentType.SCENE_BREAK, 0

    # Check for chapter headings
    for pattern in CHAPTER_PATTERNS:
        if re.match(pattern, upper):
            return ContentType.CHAPTER_HEADING, 1

    # Check for front matter headings
    for pattern in FRONT_MATTER_PATTERNS:
        if re.match(pattern, upper):
            return ContentType.FRONT_MATTER_HEADING, 2

    # Check for back matter headings
    for pattern in BACK_MATTER_PATTERNS:
        if re.match(pattern, upper):
            return ContentType.BACK_MATTER_HEADING, 2

    # Regular paragraph
    return ContentType.PARAGRAPH, 0


def parse_manuscript(content: str) -> List[ContentBlock]:
    """
    Parse manuscript content into structured content blocks.

    Args:
        content: Raw manuscript text

    Returns:
        List of ContentBlock objects
    """
    lines = content.split('\n')
    blocks: List[ContentBlock] = []
    current_paragraph_lines: List[str] = []
    blank_count = 0

    def flush_paragraph():
        """Flush accumulated paragraph lines into a block"""
        nonlocal current_paragraph_lines
        if current_paragraph_lines:
            text = ' '.join(current_paragraph_lines)
            # Clean up extra spaces
            text = re.sub(r'\s+', ' ', text).strip()
            if text:
                blocks.append(ContentBlock(ContentType.PARAGRAPH, text))
            current_paragraph_lines = []

    for line in lines:
        stripped = line.strip()
        content_type, level = detect_content_type(stripped)

        if content_type == ContentType.BLANK:
            blank_count += 1
            # Multiple blank lines might indicate a scene break or section end
            if blank_count >= 3 and current_paragraph_lines:
                flush_paragraph()
                blocks.append(ContentBlock(ContentType.SCENE_BREAK, ""))
            elif blank_count >= 1:
                # Single blank line ends current paragraph
                flush_paragraph()
            continue

        blank_count = 0

        if content_type in [ContentType.CHAPTER_HEADING, ContentType.FRONT_MATTER_HEADING,
                            ContentType.BACK_MATTER_HEADING, ContentType.SCENE_BREAK]:
            # Flush any pending paragraph
            flush_paragraph()
            blocks.append(ContentBlock(content_type, stripped, level))
        else:
            # Accumulate paragraph lines
            current_paragraph_lines.append(stripped)

    # Flush final paragraph
    flush_paragraph()

    return blocks


def clean_chapter_title(title: str) -> str:
    """
    Clean up redundant chapter titles like 'Chapter 1: Chapter One: Title'
    to just 'Chapter One: Title' or 'Chapter 1: Title'
    """
    # Pattern: "Chapter X: Chapter Word:" -> just "Chapter Word:"
    match = re.match(r'^(CHAPTER\s+\d+):\s*(CHAPTER\s+[A-Z]+[\w\s-]*:)', title, re.IGNORECASE)
    if match:
        return match.group(2).strip()

    # Pattern: "Chapter Word: Chapter X:" -> just "Chapter Word:"
    match = re.match(r'^(CHAPTER\s+[A-Z]+[\w\s-]*):\s*(CHAPTER\s+\d+:)', title, re.IGNORECASE)
    if match:
        return match.group(1).strip() + ':'

    return title


def create_professional_docx(content: str, config: DocxConfig, output_path: Path) -> Path:
    """
    Create a professionally formatted DOCX document from manuscript content.

    Args:
        content: Raw manuscript text
        config: Document configuration
        output_path: Path to save the DOCX file

    Returns:
        Path to the created DOCX file
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError("python-docx is not available. Install it with: pip install python-docx")

    logger.info(f"Creating professional DOCX: {output_path}")

    # Parse the manuscript
    blocks = parse_manuscript(content)
    logger.info(f"Parsed {len(blocks)} content blocks")

    # Create document
    doc = Document()

    # Set up page size and margins
    section = doc.sections[0]
    section.page_width = Inches(config.page_width_inches)
    section.page_height = Inches(config.page_height_inches)
    section.top_margin = Inches(config.top_margin_inches)
    section.bottom_margin = Inches(config.bottom_margin_inches)
    section.left_margin = Inches(config.left_margin_inches)
    section.right_margin = Inches(config.right_margin_inches)

    # Create custom styles
    _create_styles(doc, config)

    # Add title page if configured
    if config.include_title_page:
        _add_title_page(doc, config)

    # Add copyright page if configured
    if config.include_copyright:
        _add_copyright_page(doc, config)

    # Add dedication if configured
    if config.include_dedication and config.dedication_text:
        _add_dedication_page(doc, config)

    # Track chapter count for logging
    chapter_count = 0
    is_first_paragraph_in_chapter = False

    # Process content blocks
    for i, block in enumerate(blocks):
        if block.content_type == ContentType.CHAPTER_HEADING:
            chapter_count += 1
            # Add page break before chapter (except first)
            if chapter_count > 1 or config.include_title_page:
                doc.add_page_break()

            # Clean up redundant chapter titles
            clean_title = clean_chapter_title(block.text)

            # Add chapter heading
            para = doc.add_paragraph(clean_title, style='ChapterHeading')
            is_first_paragraph_in_chapter = True
            logger.debug(f"Added chapter: {clean_title[:50]}...")

        elif block.content_type == ContentType.FRONT_MATTER_HEADING:
            doc.add_page_break()
            para = doc.add_paragraph(block.text, style='SectionHeading')
            is_first_paragraph_in_chapter = True

        elif block.content_type == ContentType.BACK_MATTER_HEADING:
            doc.add_page_break()
            para = doc.add_paragraph(block.text, style='SectionHeading')
            is_first_paragraph_in_chapter = True

        elif block.content_type == ContentType.SCENE_BREAK:
            # Add scene break ornament
            para = doc.add_paragraph(config.scene_break_symbol, style='SceneBreak')

        elif block.content_type == ContentType.PARAGRAPH:
            if block.text:
                # Use different style for first paragraph after heading (no indent)
                if is_first_paragraph_in_chapter:
                    para = doc.add_paragraph(block.text, style='FirstParagraph')
                    is_first_paragraph_in_chapter = False
                else:
                    para = doc.add_paragraph(block.text, style='BodyText')

    # Add about author if configured
    if config.include_about_author and config.about_author_text:
        _add_about_author(doc, config)

    # Add headers and footers
    _add_headers_footers(doc, config)

    # Save document
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    logger.info(f"Created DOCX with {chapter_count} chapters: {output_path}")
    return output_path


def _create_styles(doc: Document, config: DocxConfig):
    """Create custom styles for the document"""
    styles = doc.styles

    # Chapter Heading style
    if 'ChapterHeading' not in [s.name for s in styles]:
        chapter_style = styles.add_style('ChapterHeading', WD_STYLE_TYPE.PARAGRAPH)
        chapter_style.font.name = config.heading_font
        chapter_style.font.size = Pt(config.chapter_font_size_pt)
        chapter_style.font.bold = True
        chapter_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        chapter_style.paragraph_format.space_before = Pt(72)  # 1 inch before
        chapter_style.paragraph_format.space_after = Pt(36)  # 0.5 inch after
        chapter_style.paragraph_format.keep_with_next = True

    # Section Heading style (for front/back matter)
    if 'SectionHeading' not in [s.name for s in styles]:
        section_style = styles.add_style('SectionHeading', WD_STYLE_TYPE.PARAGRAPH)
        section_style.font.name = config.heading_font
        section_style.font.size = Pt(14)
        section_style.font.bold = False
        section_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        section_style.paragraph_format.space_before = Pt(72)
        section_style.paragraph_format.space_after = Pt(24)
        # Small caps effect via font
        section_style.font.all_caps = True

    # First Paragraph style (no indent after heading)
    if 'FirstParagraph' not in [s.name for s in styles]:
        first_para_style = styles.add_style('FirstParagraph', WD_STYLE_TYPE.PARAGRAPH)
        first_para_style.font.name = config.body_font
        first_para_style.font.size = Pt(config.body_font_size_pt)
        first_para_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        first_para_style.paragraph_format.first_line_indent = Inches(0)  # No indent
        first_para_style.paragraph_format.space_after = Pt(0)
        first_para_style.paragraph_format.line_spacing = config.line_spacing

    # Body Text style
    if 'BodyText' not in [s.name for s in styles]:
        body_style = styles.add_style('BodyText', WD_STYLE_TYPE.PARAGRAPH)
        body_style.font.name = config.body_font
        body_style.font.size = Pt(config.body_font_size_pt)
        body_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        body_style.paragraph_format.first_line_indent = Inches(config.first_line_indent_inches)
        body_style.paragraph_format.space_after = Pt(0)
        body_style.paragraph_format.line_spacing = config.line_spacing

    # Scene Break style
    if 'SceneBreak' not in [s.name for s in styles]:
        break_style = styles.add_style('SceneBreak', WD_STYLE_TYPE.PARAGRAPH)
        break_style.font.name = config.body_font
        break_style.font.size = Pt(config.body_font_size_pt)
        break_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        break_style.paragraph_format.space_before = Pt(18)
        break_style.paragraph_format.space_after = Pt(18)

    # Title style
    if 'BookTitle' not in [s.name for s in styles]:
        title_style = styles.add_style('BookTitle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = config.heading_font
        title_style.font.size = Pt(config.title_font_size_pt)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_style.paragraph_format.space_after = Pt(12)

    # Subtitle style
    if 'BookSubtitle' not in [s.name for s in styles]:
        subtitle_style = styles.add_style('BookSubtitle', WD_STYLE_TYPE.PARAGRAPH)
        subtitle_style.font.name = config.heading_font
        subtitle_style.font.size = Pt(16)
        subtitle_style.font.italic = True
        subtitle_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_style.paragraph_format.space_after = Pt(24)

    # Author style
    if 'BookAuthor' not in [s.name for s in styles]:
        author_style = styles.add_style('BookAuthor', WD_STYLE_TYPE.PARAGRAPH)
        author_style.font.name = config.heading_font
        author_style.font.size = Pt(14)
        author_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        author_style.paragraph_format.space_before = Pt(36)

    # Copyright style
    if 'Copyright' not in [s.name for s in styles]:
        copyright_style = styles.add_style('Copyright', WD_STYLE_TYPE.PARAGRAPH)
        copyright_style.font.name = config.body_font
        copyright_style.font.size = Pt(10)
        copyright_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        copyright_style.paragraph_format.space_after = Pt(12)

    # Dedication style
    if 'Dedication' not in [s.name for s in styles]:
        dedication_style = styles.add_style('Dedication', WD_STYLE_TYPE.PARAGRAPH)
        dedication_style.font.name = config.body_font
        dedication_style.font.size = Pt(12)
        dedication_style.font.italic = True
        dedication_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _add_title_page(doc: Document, config: DocxConfig):
    """Add a title page to the document"""
    # Add some vertical space
    for _ in range(6):
        doc.add_paragraph()

    # Title
    doc.add_paragraph(config.title, style='BookTitle')

    # Subtitle
    if config.subtitle:
        doc.add_paragraph(config.subtitle, style='BookSubtitle')

    # Author
    if config.author:
        doc.add_paragraph(config.author, style='BookAuthor')

    # Page break after title page
    doc.add_page_break()


def _add_copyright_page(doc: Document, config: DocxConfig):
    """Add a copyright page to the document"""
    # Add some vertical space to push content down
    for _ in range(15):
        doc.add_paragraph()

    # Copyright notice
    year = config.copyright_year or "2024"
    holder = config.copyright_holder or config.author or "Author"

    doc.add_paragraph(f"Copyright © {year} {holder}", style='Copyright')
    doc.add_paragraph("All rights reserved.", style='Copyright')
    doc.add_paragraph()
    doc.add_paragraph(
        "No part of this book may be reproduced in any form or by any electronic "
        "or mechanical means, including information storage and retrieval systems, "
        "without written permission from the author, except for the use of brief "
        "quotations in a book review.",
        style='Copyright'
    )

    # Page break after copyright
    doc.add_page_break()


def _add_dedication_page(doc: Document, config: DocxConfig):
    """Add a dedication page to the document"""
    # Add vertical space
    for _ in range(8):
        doc.add_paragraph()

    doc.add_paragraph(config.dedication_text, style='Dedication')

    # Page break after dedication
    doc.add_page_break()


def _add_about_author(doc: Document, config: DocxConfig):
    """Add an About the Author section at the end"""
    doc.add_page_break()

    doc.add_paragraph("About the Author", style='SectionHeading')

    para = doc.add_paragraph(config.about_author_text, style='FirstParagraph')


def _add_headers_footers(doc: Document, config: DocxConfig):
    """Add headers and footers with page numbers"""
    section = doc.sections[0]

    # Enable different first page header/footer
    section.different_first_page_header_footer = True

    # Footer with page numbers (centered)
    footer = section.footer
    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add page number field
    _add_page_number(footer_para)

    # Header with book title (for main pages)
    header = section.header
    header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = header_para.add_run(config.title)
    run.font.size = Pt(10)
    run.font.italic = True


def _add_page_number(paragraph):
    """Add a page number field to a paragraph"""
    run = paragraph.add_run()

    # Create the field code for page number
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def generate_docx_from_manuscript(
    manuscript_path: Path,
    output_path: Path,
    config_dict: Optional[Dict[str, Any]] = None
) -> Path:
    """
    High-level function to generate a DOCX from a manuscript file.

    Args:
        manuscript_path: Path to the manuscript file
        output_path: Path for the output DOCX
        config_dict: Optional configuration dictionary

    Returns:
        Path to the generated DOCX
    """
    # Read manuscript content
    manuscript_path = Path(manuscript_path)

    if manuscript_path.suffix.lower() == '.docx':
        # Extract text from existing DOCX
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(str(manuscript_path))
            content = '\n\n'.join(p.text for p in doc.paragraphs)
        except Exception as e:
            logger.error(f"Error reading DOCX: {e}")
            raise
    else:
        # Read as text
        content = manuscript_path.read_text(encoding='utf-8', errors='ignore')

    # Build config
    config_dict = config_dict or {}
    config = DocxConfig(
        title=config_dict.get('title', manuscript_path.stem),
        author=config_dict.get('author', ''),
        subtitle=config_dict.get('subtitle', ''),
        page_width_inches=_trim_to_width(config_dict.get('trim', '6x9')),
        page_height_inches=_trim_to_height(config_dict.get('trim', '6x9')),
        top_margin_inches=config_dict.get('top_margin_in', 0.75),
        bottom_margin_inches=config_dict.get('bottom_margin_in', 0.85),
        left_margin_inches=config_dict.get('outer_margin_in', 0.75),
        right_margin_inches=config_dict.get('gutter_in', 0.7),
        body_font=_extract_font_name(config_dict.get('font_family', 'Georgia')),
        body_font_size_pt=config_dict.get('font_size_pt', 11.0),
        line_spacing=config_dict.get('line_height', 1.35),
        include_title_page=True,
        include_copyright=config_dict.get('include_copyright', True),
        copyright_year=config_dict.get('copyright_year', ''),
        copyright_holder=config_dict.get('copyright_holder', ''),
        include_dedication=config_dict.get('include_dedication', False),
        dedication_text=config_dict.get('dedication_text', ''),
        include_about_author=config_dict.get('include_about_author', False),
        about_author_text=config_dict.get('about_author_text', ''),
        scene_break_symbol=config_dict.get('scene_break', '* * *'),
    )

    return create_professional_docx(content, config, output_path)


def _trim_to_width(trim: str) -> float:
    """Convert trim size to page width in inches"""
    trim_map = {
        '5x8': 5.0,
        '5.5x8.5': 5.5,
        '6x9': 6.0,
        '8.5x11': 8.5,
    }
    return trim_map.get(trim, 6.0)


def _trim_to_height(trim: str) -> float:
    """Convert trim size to page height in inches"""
    trim_map = {
        '5x8': 8.0,
        '5.5x8.5': 8.5,
        '6x9': 9.0,
        '8.5x11': 11.0,
    }
    return trim_map.get(trim, 9.0)


def _extract_font_name(font_stack: str) -> str:
    """Extract the first font name from a CSS font stack"""
    # Remove quotes and get first font
    fonts = font_stack.split(',')
    if fonts:
        font = fonts[0].strip().strip("'\"")
        # Map common CSS fonts to Word-compatible names
        font_map = {
            'EB Garamond': 'Garamond',
            'serif': 'Georgia',
            'sans-serif': 'Arial',
        }
        return font_map.get(font, font)
    return 'Georgia'


# For testing
if __name__ == '__main__':
    import sys
    if len(sys.argv) < 2:
        print("Usage: python docx_generator.py <manuscript.txt> [output.docx]")
        sys.exit(1)

    input_file = Path(sys.argv[1])
    output_file = Path(sys.argv[2]) if len(sys.argv) > 2 else input_file.with_suffix('.docx')

    config = DocxConfig(
        title=input_file.stem.replace('_', ' ').title(),
        author="Unknown Author",
    )

    content = input_file.read_text(encoding='utf-8', errors='ignore')
    create_professional_docx(content, config, output_file)
    print(f"Created: {output_file}")
